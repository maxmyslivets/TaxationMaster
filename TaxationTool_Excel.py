import re
from collections import Counter
from tkinter.filedialog import askopenfilename

import pandas as pd
import xlwings as xw
from docx import Document as DocxDocument

from src.autocad import AutocadWorker
from src.excel import ExcelWorker
from src.parsing import Templates, Parser, Splitter
from src.validation import SearchAmbiguity


@xw.sub
def insert_word_taxation_list():
    """Вставка Word ведомости таксации"""
    file_path = askopenfilename(filetypes=[("Word files", "*.docx *.doc"), ("All files", "*.*")])
    if not file_path:
        return

    doc = DocxDocument(file_path)
    data = []

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

    df = pd.DataFrame(data)
    df.index.name = 'index'
    df.index = df.index - 1
    df.index = df.index.astype(str)

    sheet = xw.sheets['Ведомость']
    sheet['A1'].value = ['Индекс', 'Номер точки', 'Наименование', 'Количество', 'Высота', 'Толщина', 'Состояние',
                         'Кустарник', 'Неоднозначность', 'Пень', 'Наименование (пень)']
    for l in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        sheet[f'{l}:{l}'].number_format = '@'
    sheet['A2'].options(index=True, header=False).value = df


# @xw.sub
# def create_table_taxation_list():
#     """Преобразовать ведомость таксации в таблицу"""
#     selected_cells = xw.apps.active.selection
#     xw.sheets.active.tables.add(source=selected_cells, name='ВедомостьТаксации', table_style_name='TableStyleMedium9')


@xw.sub
def get_count_tree():
    """Подсчёт количества деревьев"""
    app = xw.apps.active

    selected_cells = xw.apps.active.selection
    df = pd.DataFrame(data=selected_cells.value, columns=['Количество'], dtype=str)

    def get_count(text: str) -> int:
        if re.match(Templates.DIGIT, text) or re.match(Templates.FLOAT_DIGIT, text):
            return int(float(text.strip()))
        elif re.match(Templates.TRUNKS, text):
            return 1
        else:
            return 0

    df_count = df['Количество'].apply(get_count)

    app.alert(f"Найдено деревьев: {df_count.sum()}", "Подсчёт количества деревьев")


@xw.sub
def replace_comma_to_dot():
    """Замена запятой на точку"""
    selected_cells = xw.apps.active.selection
    for cell in selected_cells:
        if not xw.sheets.active.api.Rows(cell.row).Hidden and ',' in str(cell.value):
            cell.value = str(cell.value).replace(',', '.')


@xw.sub
def replace_dot_comma_to_comma():
    """Замена точки с запятой на запятую"""
    selected_cells = xw.apps.active.selection
    for cell in selected_cells:
        if not xw.sheets.active.api.Rows(cell.row).Hidden and ';' in str(cell.value):
            cell.value = str(cell.value).replace(';', ',')


@xw.func
def get_is_shrub(name: str) -> int:
    """Определить кустарник"""
    is_shrub = Parser.get_specie(name).is_shrub
    return int(is_shrub)


@xw.sub
def get_is_shrub_sub():
    """Вставка формулы: Определить кустарник"""
    xw.apps.active.selection.formula = "=get_is_shrub([@Наименование])"


@xw.func
def check_ambiguity(number, quantity, height, diameter, is_shrub) -> int:
    """Поиск неоднозначностей в количественных характеристиках"""
    number, quantity, height, diameter = str(number), str(quantity), str(height), str(diameter)
    result = SearchAmbiguity.check_in_row_from_taxation_list(number, quantity, height, diameter, is_shrub)
    return int(not result)


@xw.sub
def check_ambiguity_sub():
    """Вставка формулы: Поиск неоднозначностей в количественных характеристиках"""
    xw.apps.active.selection.formula = ("=check_ambiguity([@Номер точки],[@Количество],[@Высота],[@Толщина],"
                                        "[@Кустарник])")


@xw.func
def identification_stump(height, diameter, is_shrub) -> int:
    """Определить пень"""
    is_stump = Parser.identification_stump(height, diameter, is_shrub)
    return int(is_stump)


@xw.sub
def identification_stump_sub():
    """Вставка формулы: Определить пень"""
    xw.apps.active.selection.formula = "=identification_stump([@Высота],[@Толщина],[@Кустарник])"


@xw.func
def insert_stump(name, is_stump) -> str:
    """Вставить пень"""
    assert "пень" not in name.lower()
    if is_stump:
        return name + " (пень)"
    else:
        return name


@xw.sub
def insert_stump_sub():
    """Вставка формулы: Вставить пень"""
    xw.apps.active.selection.formula = "=insert_stump([@Наименование],[@Пень])"


@xw.sub
def compare_numbers():
    """Сравнить наличие номеров в Excel и Autocad"""
    numbers_from_excel = ExcelWorker.get_numbers(xw.sheets['Ведомость'], 'Номер точки')
    numbers_from_acad = AutocadWorker.get_numbers('номера')
    split_numbers_from_excel = []
    for number_excel in numbers_from_excel:
        split_numbers_from_excel.extend(Splitter.number(str(number_excel)))
    split_numbers_from_acad = []
    for number_acad in numbers_from_acad:
        split_numbers_from_acad.extend(Splitter.number(str(number_acad)))
    counter_split_numbers_from_excel = Counter(split_numbers_from_excel)
    duplicates_split_numbers_from_excel = [key for key, value in counter_split_numbers_from_excel.items() if value > 1]
    counter_split_numbers_from_acad = Counter(split_numbers_from_acad)
    duplicates_split_numbers_from_acad = [key for key, value in counter_split_numbers_from_acad.items() if value > 1]
    unique_in_excel = list(set(split_numbers_from_excel) - set(split_numbers_from_acad))
    unique_in_acad = list(set(split_numbers_from_acad) - set(split_numbers_from_excel))
    new_doc = xw.books.add()
    new_doc.sheets[0]['A1'].value = ["Дубликаты Excel", "Дубликаты Autocad", "Только в Excel", "Только в Autocad"]
    i = 1
    for value in duplicates_split_numbers_from_excel:
        i += 1
        new_doc.sheets[0][f'A{i}'].value = value
    i = 1
    for value in duplicates_split_numbers_from_acad:
        i += 1
        new_doc.sheets[0][f'B{i}'].value = value
    i = 1
    for value in unique_in_excel:
        i += 1
        new_doc.sheets[0][f'C{i}'].value = value
    i = 1
    for value in unique_in_acad:
        i += 1
        new_doc.sheets[0][f'D{i}'].value = value


@xw.sub
def insert_taxation_data_from_autocad():
    """Чтение таксационных данных из топографического плана Autocad в буфер обмена"""
    topographic_plan = AutocadWorker.get_df_topographic_plan(["номера"], ["полосы"], ["контуры"], wkt_convert=True)
    sheet = xw.sheets['Автокад']
    for l in ['A', 'B', 'D']:
        sheet[f'{l}:{l}'].number_format = '@'
    sheet.range('A1').value = topographic_plan
    sheet["A1"].value = ['index']


# @xw.sub
# def create_table_taxation_plan():
#     """Преобразовать таксационные данные из топографического плана в таблицу"""
#     sheet = xw.sheets['Автокад']
#     sheet_range = sheet["A1"].expand()
#     xw.sheets.active.tables.add(source=sheet_range, name='ТаксацияАвтокад', table_style_name='TableStyleMedium9')


@xw.sub
def insert_zones_from_autocad():
    """Вставить зоны из топографического плана в таблицу"""
    zones = AutocadWorker.get_df_zones(['зоны'], wkt_convert=True)
    sheet = xw.sheets['Зоны']
    sheet.range('A1').value = zones
    sheet["A1"].value = ['index']


@xw.func
def insert_zone_objects(zone_name) -> str:
    """Вставить объекты для зоны"""
    pass


@xw.sub
def insert_zone_objects_sub():
    """Вставка формулы: Вставить объекты для зоны"""
    xw.apps.active.selection.formula = "=insert_zone_objects(Зоны!B2)"


# def main():
#     wb = xw.Book.caller()
#
#
# if __name__ == "__main__":
#     xw.Book("proj1.xlsm").set_mock_caller()
#     main()
