import re
import string
from tkinter.filedialog import askopenfilename
from docx import Document as DocxDocument

import xlwings as xw
import pandas as pd
from shapely.geometry.linestring import LineString
from shapely.geometry.multilinestring import MultiLineString
from shapely.geometry.multipolygon import MultiPolygon
from shapely.geometry.point import Point
from shapely.geometry.polygon import Polygon
from shapely.wkt import loads
from win32com.universal import com_error

from src.parsing import Splitter, Parser, Templates


class ExcelWorker:

    @staticmethod
    def index_from_cell(cell: xw.main.Range) -> str:
        sheet = xw.sheets.active
        return sheet.range(f'A{cell.row}').value

    @staticmethod
    def selected_cells(app) -> list[xw.main.Range]:
        if not app.ui.checkBox_skip_hidden_cells.isChecked():
            return xw.apps.active.selection
        else:
            selected_cells = xw.apps.active.selection
            progress = app.progress_manager.new("Сбор выделенных ячеек", len(selected_cells))
            not_hidden_cells = []
            for cell in selected_cells:
                if not xw.sheets.active.api.Rows(cell.row).Hidden:
                    not_hidden_cells.append(cell)
                progress.next()
            return not_hidden_cells

    @staticmethod
    def column_from_title(sheet: xw.Sheet) -> dict[str: str]:
        titles: list[str] = sheet.range('A1').expand('right').value
        letters: list[str] = list(string.ascii_uppercase)
        return {title: letter for title, letter in zip(titles, letters)}

    @staticmethod
    def set_text_format(sheet: xw.Sheet, num_columns: list[int]) -> None:
        letters = list(string.ascii_uppercase)
        for i in num_columns:
            sheet[f'{letters[i-1]}:{letters[i-1]}'].number_format = '@'

    @staticmethod
    def get_taxation_list(app) -> pd.DataFrame:
        """
        Получение ведомости таксации из Word

        Returns:
            pd.DataFrame: ведомость таксации
        """
        file_path = askopenfilename(filetypes=[("Word files", "*.docx *.doc"), ("All files", "*.*")])

        assert file_path is not None

        doc = DocxDocument(file_path)
        data = []
        progress_1 = app.progress_manager.new("Чтение таблиц", len(doc.tables))
        for table in doc.tables:
            progress_2 = app.progress_manager.new("Обработка строк", len(table.rows))
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                data.append(row_data)
                progress_2.next()
            progress_1.next()

        return pd.DataFrame(data)

    @staticmethod
    def get_numbers(sheet: xw.Sheet, column_name: str) -> list[str]:
        table_range = sheet['A1'].expand('table')
        data = table_range.value
        headers = data[0]
        col_index = headers.index(column_name)
        return [row[col_index] for row in data[1:]]

    @staticmethod
    def split_taxation_list_item(df: pd.DataFrame, series: pd.Series) -> list[dict]:
        """
        Получение строки ОРМ из строки таблицы Ведомости

        Args:
            df (pd.DataFrame): Датафрейм листа Автокад
            series (pd.Series): Строка таблицы

        Returns:
            list[dict]: Список словарей ОРМ для датафрейма
        """
        match_trunk = re.search(Templates.TRUNKS, series['Количество'])
        match_contour = re.search(Templates.CONTOUR, series['Количество'])
        match_line = re.search(Templates.LINE, series['Количество'])
        shapes = ExcelWorker.get_shapes_from_autocad_df(df, series['Номер точки'])
        numbers_positions, geometries = shapes['Список позиций номеров'], shapes['Список геометрии']
        if not match_contour and not match_line and not match_trunk:
            split_numbers = Splitter.number(series['Номер точки'])
            split_height = Splitter.size(series['Высота'])
            split_diameter = Splitter.size(series['Толщина'])
            split_quality = Splitter.quality(series['Состояние'])
            is_stump = Parser.identification_stump(series['Высота'], series['Толщина'], bool(series['Кустарник']))
            if len(split_numbers) == 1:
                if (not is_stump or "пень" in series['Наименование'].lower()) and len(split_quality) == 1 and len(
                        numbers_positions) == 1 and len(geometries) == 1:
                    series_dict = series.to_dict()
                    series_dict['Позиция номера'] = numbers_positions[0]
                    series_dict['Геометрия'] = geometries[0]
                    return [series_dict]
            else:
                if (not is_stump or "пень" in series['Наименование'].lower()) and len(split_quality) == 1 and len(
                        numbers_positions) == 1 and len(geometries) == 1:
                    series_dict = series.to_dict()
                    series_dict['Позиция номера'] = numbers_positions[0]
                    series_dict['Геометрия'] = geometries[0]
                    return [series_dict]
                if len(split_height) == 1:
                    split_height = split_height * int(series['Количество'])
                if len(split_diameter) == 1:
                    split_diameter = split_diameter * int(series['Количество'])
                if len(split_quality) == 1:
                    split_quality = split_quality * int(series['Количество'])
                series_data = []
                for idx in range(len(split_numbers)):
                    if "пень" not in series['Наименование'].lower():
                        is_stump = Parser.identification_stump(split_height[idx], split_diameter[idx],
                                                               bool(series['Кустарник']))
                        name = series['Наименование'] + " (пень)" if is_stump else series['Наименование']
                    else:
                        name = series['Наименование']
                    series_data.append({
                        'Номер точки': split_numbers[idx],
                        'Наименование': name,
                        'Количество': 1,
                        'Высота': split_height[idx],
                        'Толщина': split_diameter[idx],
                        'Состояние': split_quality[idx],
                        'Кустарник': series['Кустарник'],
                        'Позиция номера': numbers_positions[idx],
                        'Геометрия': geometries[idx]
                    })
                return series_data
        else:
            series_dict = series.to_dict()
            series_dict['Позиция номера'] = numbers_positions[0]
            series_dict['Геометрия'] = geometries[0]
            return [series_dict]

    @staticmethod
    def get_shapes_from_autocad_df(df: pd.DataFrame, number: str) -> dict:
        """
        Получение геометрии из датафрейма "Автокад" по номеру точки
        Args:
            df (pd.DataFrame): Датафрейм "Автокад"
            number (str): Номер точки

        Returns:
            dict: Словарь двух колонок для датафрейма со списками позиций номеров и геометрий
        """
        number_positions, geometries = [], []
        split_numbers = Splitter.number(number)
        df = df.set_index('Разделенный номер')
        for split_number in split_numbers:
            shapes = df.loc[split_number][['Позиция номера', 'Геометрия']].to_dict()
            number_positions.append(shapes['Позиция номера'])
            geometries.append(shapes['Геометрия'])
        return {'Список позиций номеров': number_positions, 'Список геометрии': geometries}

    @staticmethod
    def get_taxation_list_orm(wkt_convert: bool = False, app=None) -> pd.DataFrame:
        """
        Получение датафрейма "Ведомости ОРМ" из таблиц "Ведомость" и "Автокад"

        Args:
            wkt_convert (bool): Преобразование геометрии в wkt
            app (Progress): Прогресс бар

        Returns:
            pd.DataFrame: Датафрейм ведомости ОРМ
        """
        progress_1 = app.progress_manager.new("Получение датафрейма", 2)
        sheet_taxation_list = xw.sheets['Ведомость']
        taxation_list_df = sheet_taxation_list.range('A1').expand().options(pd.DataFrame, header=1).value
        taxation_list_df = taxation_list_df[
            ['Номер точки', 'Наименование', 'Количество', 'Высота', 'Толщина', 'Состояние', 'Кустарник']]

        sheet_autocad = xw.sheets['Автокад']
        autocad_df = sheet_autocad.range('A1').expand().options(pd.DataFrame, header=1, index=False).value
        autocad_df['Позиция номера'] = autocad_df['Позиция номера'].apply(lambda x: loads(x))
        autocad_df['Геометрия'] = autocad_df['Геометрия'].apply(lambda x: loads(x))

        assert autocad_df['Разделенный номер'].is_unique

        progress_1.set_value(1)
        progress_2 = app.progress_manager.new("Разделение ОРМ", len(taxation_list_df))

        taxation_list_orm = []
        for _, series in taxation_list_df.iterrows():
            taxation_list_orm.extend(ExcelWorker.split_taxation_list_item(autocad_df, series))
            progress_2.next()

        taxation_list_orm_df = pd.DataFrame(taxation_list_orm)

        progress_1.set_value(2)

        if not wkt_convert:
            return taxation_list_orm_df
        else:
            taxation_list_orm_df['Позиция номера'] = taxation_list_orm_df['Позиция номера'].apply(lambda geom: geom.wkt)
            taxation_list_orm_df['Геометрия'] = taxation_list_orm_df['Геометрия'].apply(lambda geom: geom.wkt)
            return taxation_list_orm_df

    @staticmethod
    def get_size_text(geometry: Point | Polygon | MultiPolygon | LineString | MultiLineString, quantity: str) -> str:
        if isinstance(geometry, Polygon) or isinstance(geometry, MultiPolygon):
            return f"{round(geometry.area, 1)} м.кв."
        elif isinstance(geometry, LineString) or isinstance(geometry, MultiLineString):
            return f"{round(geometry.length, 1)} м.п."
        else:
            try:
                assert int(quantity) % 1 == 0
                return str(int(float(quantity)))
            except:
                return quantity

    @staticmethod
    def get_objects_from_zone(zone_result: str, wkt_convert: bool = False, app=None) -> pd.DataFrame:
        """
        Получить датафрейм объектов, входящих в указанную зону

        Args:
            zone_result (str): Название зоны
            wkt_convert (bool): Преобразование геометрии в wkt
            app: Прогресс бар

        Returns:
            pd.DataFrame: Датафрейм объектов, входящих в зону
        """
        progress_1 = app.progress_manager.new("Получение датафрейма", 100)
        sheet_autocad = xw.sheets['Ведомость ОРМ']
        taxation_list_orm_df = sheet_autocad.range('A1').expand().options(pd.DataFrame, header=1, index=False).value
        taxation_list_orm_df['Позиция номера'] = taxation_list_orm_df['Позиция номера'].apply(lambda x: loads(x))
        taxation_list_orm_df['Геометрия'] = taxation_list_orm_df['Геометрия'].apply(lambda x: loads(x))

        sheet_zones = xw.sheets['Зоны']
        zones_df = sheet_zones.range('A1').expand().options(pd.DataFrame, header=1, index=False).value

        zone_names = zones_df['Наименование'].tolist()
        zone_names.remove(zone_result)

        used_split_numbers_df = pd.DataFrame(columns=['Исх.номер', 'Список геометрии'])

        progress_1.set_value(30)
        progress_2 = app.progress_manager.new("Фильтрация объектов", len(zone_names))
        for zone_name in zone_names:
            try:
                sheet_zone = xw.sheets[zone_name]
            except com_error:
                progress_2.next()
                continue
            if sheet_zone.range('A1').value is None:
                progress_2.next()
                continue
            _used_split_numbers_df = sheet_zone.range('A1').expand().options(pd.DataFrame, header=1, index=False).value[
                ['Исх.номер', 'Геометрия']]
            _used_split_numbers_df['Геометрия'] = _used_split_numbers_df['Геометрия'].apply(lambda x: loads(x))
            used_split_numbers_df = pd.concat([used_split_numbers_df, _used_split_numbers_df])
            progress_2.next()

        if len(used_split_numbers_df) != 0:
            used_split_numbers_df['type'] = used_split_numbers_df['Геометрия'].apply(lambda x: isinstance(x, Point))
            used_split_numbers = used_split_numbers_df[used_split_numbers_df['type']]['Исх.номер'].tolist()
        else:
            used_split_numbers = []

        taxation_list_orm_df_not_used = taxation_list_orm_df[~taxation_list_orm_df['Номер точки'].isin(used_split_numbers)]

        zone_shape = loads(zones_df[zones_df['Наименование'] == zone_result]['Геометрия'].tolist()[0])

        progress_1.set_value(60)
        progress_2 = app.progress_manager.new("Поиск пересечений", len(zone_names))

        intersections_shapes = []
        for _, series in taxation_list_orm_df_not_used.iterrows():
            geometry = series['Геометрия']
            intersection = zone_shape.intersection(geometry)
            if intersection:
                if isinstance(intersection, Point) or isinstance(intersection, Polygon) or isinstance(intersection,
                                                                                                      LineString):
                    intersections_shapes.append({
                        'Исх.номер': series['Номер точки'],
                        'Номер': None,
                        'Наименование': series['Наименование'],
                        'Количество': ExcelWorker.get_size_text(intersection, series['Количество']),
                        'Высота': series['Высота'],
                        'Толщина': series['Толщина'],
                        'Состояние': series['Состояние'],
                        'Кустарник': series['Кустарник'],
                        'Позиция номера': intersection.centroid,
                        'Геометрия': intersection
                    })
                elif isinstance(intersection, MultiPolygon) or isinstance(intersection, MultiLineString):
                    number_suffix = 0
                    for intersection_part in intersection.geoms:
                        number_suffix += 1
                        intersections_shapes.append({
                            'Исх.номер': series['Номер точки'] + '_' + str(number_suffix),
                            'Номер': None,
                            'Наименование': series['Наименование'],
                            'Количество': ExcelWorker.get_size_text(intersection_part, series['Количество']),
                            'Высота': series['Высота'],
                            'Толщина': series['Толщина'],
                            'Состояние': series['Состояние'],
                            'Кустарник': series['Кустарник'],
                            'Позиция номера': intersection_part.centroid,
                            'Геометрия': intersection_part
                        })
            progress_2.next()

        intersections_shapes_df = pd.DataFrame(intersections_shapes)
        progress_1.set_value(100)

        if not wkt_convert:
            return intersections_shapes_df
        else:
            intersections_shapes_df.index = intersections_shapes_df.index.astype(str)
            intersections_shapes_df['Позиция номера'] = intersections_shapes_df['Позиция номера'].apply(
                lambda geom: geom.wkt if geom else None)
            intersections_shapes_df['Геометрия'] = intersections_shapes_df['Геометрия'].apply(
                lambda geom: geom.wkt if geom else None)
            return intersections_shapes_df

    @staticmethod
    def generate_numeration_from_zone(wkt_convert=True, app=None) -> pd.DataFrame:
        """
        Нумерует
        Args:
            wkt_convert:
            app:

        Returns:

        """
        # TODO: Добавить проверку имеется ли на листе current_number
        current_number = app.ui.lineEdit_start_number.text()
        start_numeration = int(app.ui.lineEdit_start_numeration.text())
        sheet = xw.sheets.active
        zone_df = sheet.range('A1').expand().options(pd.DataFrame, header=1, index=False).value
        zone_df['Геометрия'] = zone_df['Геометрия'].apply(loads)
        zone_df['Позиция номера'] = zone_df['Позиция номера'].apply(loads)
        # TODO: Добавить проверку имеется ли на других листах номера из
        #  list(range(start_numeration, start_numeration + len(zone_df['Исх.номер'])))

        current_point = zone_df['Позиция номера'][zone_df['Исх.номер'] == current_number].iloc[0]
        points = zone_df['Позиция номера'].tolist()
        progress = app.progress_manager.new("Сортировка по ближайшим", len(points))
        visited = []
        while len(visited) < len(points):
            distances = [(idx, current_point.distance(pt)) for idx, pt in enumerate(points) if idx not in visited]
            nearest_idx = min(distances, key=lambda x: x[1])[0]
            visited.append(nearest_idx)
            current_point = points[nearest_idx]
            progress.next()
        zone_df['Очередь'] = [visited.index(idx) + 1 for idx in range(len(points))]

        zone_df['Номер'] = zone_df['Очередь'].apply(lambda x: start_numeration + x - 1)

        zone_df = zone_df.drop(columns=['Индекс', 'Очередь'])

        if not wkt_convert:
            return zone_df
        else:
            zone_df.index = zone_df.index.astype(str)
            zone_df['Позиция номера'] = zone_df['Позиция номера'].apply(
                lambda geom: geom.wkt if geom else None)
            zone_df['Геометрия'] = zone_df['Геометрия'].apply(
                lambda geom: geom.wkt if geom else None)
            return zone_df

    @staticmethod
    def get_zip_numbers(app=None) -> pd.DataFrame:
        """
        Получает датафрейм с связанными номерами и их позициями
        Args:
            app:
        Returns:
            pd.DataFrame
        """
        # FIXME: Добавить колонку действия
        sheet = xw.sheets.active
        zone_df = sheet.range('A1').expand().options(pd.DataFrame, header=1, index=False).value
        zone_df['Позиция номера'] = zone_df['Позиция номера'].apply(loads)
        zone_df = zone_df[['Номер', 'Позиция номера']]

        progress_1 = app.progress_manager.new("Объединение по позициям", len(zone_df))

        zip_numbers = {}
        for _, series in zone_df.iterrows():
            if series['Позиция номера'] not in zip_numbers:
                zip_numbers[series['Позиция номера']] = [series['Номер']]
            else:
                zip_numbers[series['Позиция номера']].append(series['Номер'])
            progress_1.next()

        progress_2 = app.progress_manager.new("Сжатие номеров", len(zip_numbers))

        def compress_ranges(numbers):
            numbers = sorted(map(int, numbers))
            result = []
            start = numbers[0]
            end = start

            for i in range(1, len(numbers)):
                if numbers[i] == end + 1:
                    end = numbers[i]
                else:
                    if end - start == 1:
                        result.append(f"{start},{end}")
                    else:
                        result.append(f"{start}-{end}" if start != end else str(start))
                    start = end = numbers[i]

            if end - start == 1:
                result.append(f"{start},{end}")
            else:
                result.append(f"{start}-{end}" if start != end else str(start))

            return ",".join(result)

        zip_numbers_text = {}
        for k, v in zip_numbers.items():
            zip_numbers_text[k] = compress_ranges(v)
            progress_2.next()

        zip_numbers_text_for_df = [{'Номер': v, 'Позиция номера': k} for k, v in zip_numbers_text.items()]

        return pd.DataFrame(zip_numbers_text_for_df)
